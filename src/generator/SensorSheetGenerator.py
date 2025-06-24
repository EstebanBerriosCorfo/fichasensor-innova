import os
from datetime import datetime
from docx import Document
from src.utils.TextUtils import TextUtils
from src.utils.DateUtils import DateUtils

class SensorSheetGenerator:
    def __init__(self, template_path):
        """
        :param template_path: Ruta absoluta del archivo de plantilla Word
        """
        self.template_path = template_path

    def generate(self, meeting, project, output_dir, idx=1, meetings_list=None):
        """
        Genera el registro individual de visita/reunión.
        :param meeting: dict con los datos de una reunión
        :param project: dict con los datos generales del proyecto
        :param output_dir: carpeta de salida
        :param idx: opcional, para distinguir archivos si hay varias reuniones en la misma fecha
        :param meetings_list: opcional, lista completa de reuniones (para extraer el último ejecutivo si es necesario)
        :return: ruta del archivo generado
        """
        doc = Document(self.template_path)

        # --- Datos a extraer del meeting y project ---
        fecha_reunion = TextUtils.to_safe_str(meeting.get("Fecha de reunión", ""))
        codigo = TextUtils.to_safe_str(meeting.get("Código proyecto", ""))
        nombre_proyecto = TextUtils.to_safe_str(project.get("Nombre Proyecto") or project.get("Nombre proyecto") or "")
        beneficiario = TextUtils.to_safe_str(project.get("Nombre Beneficiario") or project.get("Beneficiario") or "")
        id_ficha = TextUtils.to_safe_str(meeting.get("Id",""))
        
        enfoque_reunion = (meeting.get("Enfoque reunión", "") or "").strip().lower()
        if enfoque_reunion == "ambas":
            actividad_lugar = "Enfoque administrativo y negocios o mentoring"
        else:
            actividad_lugar = TextUtils.to_safe_str(meeting.get("Enfoque reunión", ""))
        representante_legal = TextUtils.to_safe_str(project.get("Representante Legal") or project.get("Representante legal") or "")

        # --- Lógica para el ejecutivo ---
        ejecutivo = TextUtils.to_safe_str(
            project.get("Ejecutivo Técnico") or project.get("Nombre Ejecutivo") or ""
        )
        if not ejecutivo and meetings_list and len(meetings_list) > 0:
            ejecutivo = TextUtils.to_safe_str(meetings_list[-1].get("Nombre", ""))
        if not ejecutivo:
            ejecutivo = TextUtils.to_safe_str(meeting.get("Nombre", ""))


        #--- Lógica para observaciones y requerimientos ---

        # Limpia la celda de objetivos
        observaciones = meeting.get("Observación general", "No se registraron objetivos.")

        requerimiento_flag = meeting.get("Relativo a la ficha sensor, el/la beneficiaria hizo alguna solicitud adicional?", "")
        requerimientos = meeting.get("Requerimientos", "")

        # --- Lógica para el bloque de requerimientos ---
        if str(requerimiento_flag).strip().lower() == "si":
            requerimiento_text = "Se definieron comentarios adicionales en la reunión."
            if requerimientos and str(requerimientos).strip().lower() not in ("nan", ""):
                requerimiento_text += f"\n{TextUtils.to_safe_str(requerimientos)}"
        elif str(requerimiento_flag).strip().lower() == "no":
            requerimiento_text = "No se definieron comentarios adicionales en la reunión."
        else:
            requerimiento_text = "No se definieron comentarios adicionales en la reunión."

        # --- Concatenar todo ---
        bloques = []
        if observaciones and str(observaciones).strip().lower() not in ("nan", ""):
            bloques.append(TextUtils.to_safe_str(observaciones).strip())
        if requerimiento_text:
            bloques.append(requerimiento_text)
        texto_objetivos = "\n".join(bloques)

        # --- Lógica para compromisos ---
        compromiso_flag = (meeting.get("¿Se establecieron compromisos durante la reunión", "") or "").strip()
        detalle_compromiso = TextUtils.to_safe_str(meeting.get("Detalle el(los) compromiso(s)", ""))
        fecha_compromiso = DateUtils.format_date(meeting.get("Fecha comprometida", ""))
        responsable = TextUtils.to_safe_str(meeting.get("Responsable", ""))


        # === LLENADO DE LA PLANTILLA ===

        # Tabla 0: FECHA
        tabla_fecha = doc.tables[0]
        # Por lo general, la fecha va en la segunda celda de la primera fila
        tabla_fecha.rows[0].cells[1].text = fecha_reunion

        # Tabla 1: ACTIVIDAD
        tabla_actividad = doc.tables[1]

        tipo = (meeting.get("Tipo de reunión", "") or "").strip().lower()

        # Limpia todas las celdas (excepto los encabezados) en la primera fila
        for idx in [1, 3, 5]:
            tabla_actividad.rows[0].cells[idx].text = ""

        # Decide en qué posición poner la "X"
        if "visita a terreno" in tipo or "visita_terreno" in tipo:
            tabla_actividad.rows[0].cells[1].text = "X"
        elif ("reunión_corfo" in tipo or "reunión corfo" in tipo
            or "reunión_virtual" in tipo or "reunión virtual" in tipo):
            tabla_actividad.rows[0].cells[3].text = "X"
        else:
            tabla_actividad.rows[0].cells[5].text = "X"


        # Tabla 2: RESUMEN
        tabla_resumen = doc.tables[2]
        for row in tabla_resumen.rows:
            label = row.cells[0].text.strip().upper()
            if "CÓDIGO PROYECTO" in label:
                row.cells[1].text = codigo
            elif "NOMBRE PROYECTO" in label:
                row.cells[1].text = nombre_proyecto
            elif "BENEFICIARIO" in label:
                row.cells[1].text = beneficiario
            elif "ACTIVIDAD / LUGAR" in label:
                row.cells[1].text = actividad_lugar
            elif "EJECUTIVO" in label:
                row.cells[1].text = ejecutivo

        # Tabla 3: OBJETIVOS
        tabla_objetivos = doc.tables[3]
        # --- Insertar en la celda correspondiente ---
        tabla_objetivos.rows[1].cells[0].text = texto_objetivos


        # Tabla 4: COMPROMISOS
        tabla_compromisos = doc.tables[4]
        # Lógica principal
        if "no se registraron compromisos en esta reunión" in compromiso_flag.lower():
            tabla_compromisos.rows[1].cells[0].text = "No se registraron compromisos en esta reunión"
            tabla_compromisos.rows[1].cells[1].text = "No aplica"
            tabla_compromisos.rows[1].cells[2].text = "No se registraron responsables"
        elif compromiso_flag.lower() == "si":
            tabla_compromisos.rows[1].cells[0].text = detalle_compromiso
            tabla_compromisos.rows[1].cells[1].text = fecha_compromiso
            tabla_compromisos.rows[1].cells[2].text = responsable
        else:
            # Si hay otra respuesta o está vacío, puedes decidir cómo actuar, por ejemplo dejar vacío o marcar "No aplica"
            tabla_compromisos.rows[1].cells[0].text = "No se registraron compromisos en esta reunión"
            tabla_compromisos.rows[1].cells[1].text = "No aplica"
            tabla_compromisos.rows[1].cells[2].text = "No se registraron responsables"

        # Tabla 5: FIRMAS
        tabla_firmas = doc.tables[5]
        # Inserta los datos en la primera fila de firmas
        tabla_firmas.rows[1].cells[0].text = representante_legal
        tabla_firmas.rows[1].cells[1].text = beneficiario
        tabla_firmas.rows[1].cells[2].text = "" 

        # === Genera el nombre del archivo y guarda ===
        safe_fecha = fecha_reunion.replace("/", "-").replace(" ", "_")
        file_name = f"REGISTRO_FICHA_SENSOR_{codigo}_{safe_fecha}_{id_ficha}.docx"
        output_path = os.path.join(output_dir, file_name)
        doc.save(output_path)
        print(f"✅ Registro individual generado en: {output_path}")
        return output_path
