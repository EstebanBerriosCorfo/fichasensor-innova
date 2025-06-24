import os
from datetime import datetime
from docx import Document
from src.utils.FileUtils import FileUtils
from src.utils.TextUtils import TextUtils
from src.utils.DateUtils import DateUtils

class BitacoraGenerator:
    def __init__(self, template_path):
        self.template_path = template_path

    def generate(self, project_info, output_dir, codigo_proyecto=None):
        doc = Document(self.template_path)
        data = project_info.to_dict() if hasattr(project_info, "to_dict") else project_info

        cabecera = data.get("project", {})
        codigo = TextUtils.to_safe_str(
            codigo_proyecto
            or cabecera.get("Código Proyecto")
            or cabecera.get("Código proyecto")
            or "SIN_CODIGO"
        )
        nombre_proyecto = TextUtils.to_safe_str(
            cabecera.get("Nombre Proyecto") or cabecera.get("Nombre proyecto") or ""
        )
        beneficiario = TextUtils.to_safe_str(
            cabecera.get("Nombre Beneficiario") or cabecera.get("Beneficiario") or ""
        )
        representante_legal = TextUtils.to_safe_str(
            cabecera.get("Representante Legal") or cabecera.get("Representante legal") or ""
        )
        ejecutivo = TextUtils.to_safe_str(
            cabecera.get("Ejecutivo Técnico") or cabecera.get("Nombre Ejecutivo") or ""
        )
        meetings = data.get("meetings", [])
        if not ejecutivo and meetings:
            ejecutivo = TextUtils.to_safe_str(meetings[-1].get("Nombre", ""))

        # === ORDENAR MEETINGS POR FECHA (ascendente, más lejana a más cercana) ===
        def fecha_key(meeting):
            # Intenta convertir a datetime, si falla usa 1900-01-01 como fallback
            valor = meeting.get("Fecha de reunión", "")
            try:
                # Primero intenta tu formateador utilitario, luego iso, luego fallback
                dt = None
                try:
                    # Si viene como datetime
                    if isinstance(valor, datetime):
                        return valor
                    # Si es string, trata varios formatos comunes
                    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d", "%d-%m-%Y", "%d/%m/%Y"):
                        try:
                            return datetime.strptime(str(valor), fmt)
                        except Exception:
                            continue
                    # Si nada resultó, usa DateUtils si lo prefieres
                    return datetime.strptime(DateUtils.format_date(valor), "%d-%m-%Y")
                except Exception:
                    return datetime(1900, 1, 1)
            except Exception:
                return datetime(1900, 1, 1)

        meetings_sorted = sorted(meetings, key=fecha_key)

        # Asignar valores cabecera...
        cabecera_table = doc.tables[0]
        for row in cabecera_table.rows:
            etiqueta = row.cells[0].text.strip().upper()
            if "NOMBRE DEL PROYECTO" in etiqueta:
                row.cells[1].text = nombre_proyecto
            elif "CÓDIGO PROYECTO" in etiqueta:
                row.cells[1].text = codigo
            elif "EJECUTIVO TÉCNICO" in etiqueta or "EJECUTIVO" in etiqueta:
                row.cells[1].text = ejecutivo
            elif "BENEFICIARIO" in etiqueta:
                row.cells[1].text = beneficiario

        # === DETALLE DE REUNIONES ORDENADAS ===
        acciones_table = doc.tables[1]
        while len(acciones_table.rows) > 1:
            acciones_table._tbl.remove(acciones_table.rows[1]._tr)

        for meeting in meetings_sorted:
            row = acciones_table.add_row().cells
            fecha_reunion = meeting.get("Fecha de reunión", "")
            id_ficha = TextUtils.to_safe_str(meeting.get("Id", ""))
            safe_fecha = DateUtils.format_date(fecha_reunion).replace("/", "-").replace(" ", "_")
            row[0].text = TextUtils.to_safe_str(fecha_reunion)
            row[1].text = TextUtils.to_safe_str(meeting.get("Tipo de reunión", ""))
            row[2].text = representante_legal

            # Observaciones y lógicas asociadas
            observaciones = meeting.get("Observación general", "No se registraron objetivos.")
            compromiso_flag = meeting.get("¿Se establecieron compromisos durante la reunión", "")
            detalles = meeting.get("Detalle el(los) compromiso(s)", "")
            requerimiento_flag = meeting.get("Relativo a la ficha sensor, el/la beneficiaria hizo alguna solicitud adicional?", "")
            requerimientos = meeting.get("Requerimientos", "")

            # Compromisos
            if str(compromiso_flag).strip().lower() == "si":
                compromiso_text = "Se definieron compromisos en la reunión."
                if detalles and str(detalles).strip().lower() not in ("nan", ""):
                    compromiso_text += f"\n{TextUtils.to_safe_str(detalles)}"
            elif str(compromiso_flag).strip().lower() == "no":
                compromiso_text = "No se definieron compromisos."
            else:
                compromiso_text = "No se definieron compromisos."

            # Requerimientos
            if str(requerimiento_flag).strip().lower() == "si":
                requerimiento_text = "Se definieron comentarios adicionales en la reunión."
                if requerimientos and str(requerimientos).strip().lower() not in ("nan", ""):
                    requerimiento_text += f"\n{TextUtils.to_safe_str(requerimientos)}"
            elif str(requerimiento_flag).strip().lower() == "no":
                requerimiento_text = "No se definieron comentarios adicionales en la reunión."
            else:
                requerimiento_text = "No se definieron comentarios adicionales en la reunión."

            # Concatenar todo en observaciones
            bloques = []
            if observaciones and str(observaciones).strip().lower() not in ("nan", ""):
                bloques.append(TextUtils.to_safe_str(observaciones).strip())
            if compromiso_text:
                bloques.append(compromiso_text)
            if requerimiento_text:
                bloques.append(requerimiento_text)
            texto_obs = "\n".join(bloques)
            row[3].text = texto_obs

            # Medio de verificación
            row[4].text = (
                f'REGISTRO_FICHA_SENSOR_{codigo}_{safe_fecha}_{id_ficha}.docx'
            )

        file_name = f"BITACORA_FICHA_SENSOR_{codigo}_{datetime.now().strftime('%Y%m%d')}.docx"
        output_path = os.path.join(output_dir, file_name)
        doc.save(output_path)
        print(f"✅ Bitácora generada en: {output_path}")
        return output_path